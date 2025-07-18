import re
import sys
import json
from docx import Document
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Length, Pt,Cm,Inches
from docx.styles import style
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


page_Height = 0
title_Height = 0
header_Height = 0
# isinstance =  false


# Define the Files directory path
FILES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Files')

# Create Files directory if it doesn't exist
if not os.path.exists(FILES_DIR):
    os.makedirs(FILES_DIR)



filenumber = 0
date =  "DD-MM-YYYY"

def create_document():
    doc = Document()
    # set default font for the entire document
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    # set default paragraph format
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # set orientation to landscape
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    
    # Set margins
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

    # Calculate usable page dimensions
    usable_page_height = section.page_height.cm - (section.top_margin.cm + section.bottom_margin.cm)
    usable_page_height = usable_page_height * 0.393701 #convert cm to inches
    # print(usable_page_height)
    return doc

def create_header(doc):
    # Add header for the document
    header = doc.add_table(rows=1, cols=3)
    header_row = header.rows[0]
    
    # First cell
    header_row.cells[0].paragraphs[0].clear()  # Clear existing content
    run = header_row.cells[0].paragraphs[0].add_run("azertyuiop")
    run.font.size = Pt(13)
    font_size_1 = run.font.size.pt
    
    # Second cell
    header_row.cells[1].paragraphs[0].clear()  # Clear existing content
    run = header_row.cells[1].paragraphs[0].add_run("                         ")
    run.font.size = Pt(13)
    font_size_2 = run.font.size.pt
    
    # Third cell
    third_cell = header_row.cells[2]
    third_cell.paragraphs[0].clear()  # Clear existing content
    
    # Add each line as a separate run
    text_lines = [
        "XXXXXXXXXXXXXXXXXXXXX",
        "YYYYYYYYYYYYYYYYYYYYYYY",
        "TTTTTTTTTTTTTTTTTTTTTTTTT",
        f"A CITY, le " , 
        f"N° XXXX"  
    ]

    font_size_3 = 13 # Font size for third cell 
    line_height_3 = font_size_3 * 1.15 #Assuming 1.15x line spacing
    num_lines_3 = len(text_lines) # total lines in the third cell
    
    for i, line in enumerate(text_lines):
        run = third_cell.paragraphs[0].add_run(line)
        run.font.size = Pt(13)  # Sin this line it was laike this :  run.font.size = Pt(font_size_3) but i modify it to 13
        # Add underlining for the last two lines
        if i >= len(text_lines) - 2:
            run.font.underline = True
        if i < len(text_lines) - 1:
            run.add_break()
    # Center the text in the third cell
    third_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # !!!!! i stoped here N.B : i have to find a way to calculate the height of the header
    # Calculate the height of the header section
    # Assuming the tallest content determines the row height
    line_height_1 = font_size_1 * 1.15  # Line height for first cell
    line_height_2 = font_size_2 * 1.15  # Line height for second cell
    height_cell_1 = line_height_1  # Single line in first cell
    height_cell_2 = line_height_2  # Single line in second cell
    height_cell_3 = num_lines_3 * line_height_3  # Multi-line height in third cell
    header_Height = max(height_cell_1, height_cell_2, height_cell_3) /72
    # print(f"Height Cell 1: {height_cell_1 / 72} inches")
    # print(f"Height Cell 2: {height_cell_2 / 72} inches")
    # print(f"Height Cell 3: {height_cell_3 / 72} inches")
    # print(f"Final Estimated Header Height: {header_Height} inches")
    # print(header_Height)

    return header

def create_title_section(doc):
    # Add Body Content
    body_title = doc.add_paragraph("F I C H E")
    body_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    body_subtitle = doc.add_paragraph(
        "À\n L'Attention de SDFGHJKGFDUYTRTYUIUYTYUYTR,\n"
        "CCCCCCCCCCCCVBBBBBBBBBBBBBBBBBBB\n"
    )
    body_subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Default font size in points
    font_size_title = body_title.style.font.size.pt if body_title.style.font.size else 12  # Default to 12pt
    font_size_subtitle = body_subtitle.style.font.size.pt if body_subtitle.style.font.size else 12

    # Count the number of lines in each paragraph
    lines_title = body_title.text.count('\n') + 1  # +1 for the line itself
    lines_subtitle = body_subtitle.text.count('\n') + 1

    # Calculate height in points (font size * number of lines)
    height_title = font_size_title * lines_title
    height_subtitle = font_size_subtitle * lines_subtitle

    # Total height in points
    total_height_points = height_title + height_subtitle

    # Convert points to inches (1 point = 1/72 inch)
    title_section_height = total_height_points / 72
    # print(title_section_height)

    return body_title, body_subtitle

def create_object_section(doc):
    # First paragraph
    body_content = doc.add_paragraph()
    content_run = body_content.add_run("\t\tO B J E T : ")  # Adding spaces between letters manually
    content_run.bold = True
    content_run.font.size = Pt(12)  # Set font size
    font_size_1 = content_run.font.size.pt 

    content_run = body_content.add_run("Lorem ipsum dolor, sit amet consectetur adipisicing elit\n")
    content_run.font.size = Pt(12)  # Set font size

    # Second paragraph
    body_text = (
        "Lorem ipsum dolor, sit amet consectetur adipisicing elit. Ex labore voluptatibus "
        "sapiente magni unde beatae id provident veritatis a mollitia. Similique tempora saepe ad "
        "nam! Dolorem eos deserunt quo velit!"
    )
    body_paragraph = doc.add_paragraph(f"\t{body_text}")
    body_paragraph.style.font.size = Pt(12)  # Set font size
    font_size_2 = body_paragraph.style.font.size.pt

    # Calculate height for first paragraph
    line_height_1 = font_size_1 * 1.15  # Assuming 1.15x line spacing
    num_lines_1 = 1  # Single-line content
    height_paragraph_1 = num_lines_1 * line_height_1

    # Calculate height for second paragraph
    line_height_2 = font_size_2 * 1.15  # Assuming 1.15x line spacing
    num_lines_2 = (len(body_text) // 70) + 1  # Approximate lines (70 chars per line)
    height_paragraph_2 = num_lines_2 * line_height_2

    # Total height in points (convert to inches)
    Object_section_height = (height_paragraph_1 + height_paragraph_2) / 72  # Convert points to inches

    # print(Object_section_height)
# 
    return body_content

def create_table_section(doc, data):
    #define table header 
    table_header = [
        "Unité", "Matériel", "Date d'attribution", "Date de prise en charge", "Technicien",
        "Constat Global", "Diagnostic Technique", "Action Entreprise", "Actions Proposées"
    ]

    #create table with header 
    table = doc.add_table(rows=1, cols=len(table_header), style="Table Grid")

    # Set the width for the last column (Action Proposées) to 3 inches
    last_column = table.columns[-1]
    for cell in last_column.cells:
        cell.width = Inches(3) 
    
    # Add data rows for table after checking if data exists and is not empty
    length = len(data["Material_Data"]) 

    # Format header row with smaller font for multiple rows
    header_row = table.rows[0]
    header_font_size = 12 if length > 1 else 13  # Smaller font for multiple rows
    
    for i, text in enumerate(table_header):
        cell = header_row.cells[i]
        cell.text = text 
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(header_font_size)

    # Calculate row height based on number of rows
    row_heights = {
        1: 2.35,  # Original height for single row
        2: 1.2,   # Adjusted height for two rows
        3: 0.8,   # Adjusted height for three rows
        4: 0.6    # Adjusted height for four rows
    }
    
    row_height = row_heights.get(length, 0.6)  # Default to smallest height if length > 4

    # Add data rows based on length
    for idx in range(length):
        Material_Data = data.get('Material_Data', [])[idx]
        row = table.add_row()
        
        # Set row height
        row.height = Inches(row_height)
        
        # Fill row data
        row.cells[0].text = Material_Data.get('unit', '').upper()
        row.cells[1].text = Material_Data.get('material', '').upper()
        row.cells[2].text = Material_Data.get('attributionDate', '')
        row.cells[3].text = Material_Data.get('entryDate', '')
        row.cells[4].text = '\n'.join(Material_Data.get('technicians', [])).upper()
        row.cells[5].text = Material_Data.get('globalStatus', '')
        row.cells[6].text = Material_Data.get('technicalDiagnostic', '')
        row.cells[7].text = Material_Data.get('actionTaken', '')
        row.cells[8].text = Material_Data.get('proposedActions', '')

        # Set cell properties
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Adjust font size based on number of rows
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10 if length > 1 else 11)

    # Merge cells with same content (skip header row)
    if length > 1:  # Only attempt merge if there's more than one data row
        for col_idx in range(len(table_header)):
            row_idx = 1  # Start from first data row
            while row_idx < len(table.rows) - 1:  # -1 to avoid checking last row unnecessarily
                current_cell = table.rows[row_idx].cells[col_idx]
                next_cell = table.rows[row_idx + 1].cells[col_idx]
                
                # If cells have the same non-empty content
                if current_cell.text == next_cell.text and current_cell.text.strip():
                    # Merge the cells & Empty the next cell
                    next_cell.text = ''
                    current_cell.merge(next_cell)
                    row_idx += 2  # Skip the merged cell
                else:
                    row_idx += 1  # Move to next row

    return table

def create_footer(doc, data):
    # Add a paragraph break before the footer
    doc.add_paragraph("\n")

    assigned_persons = data.get('assignedPersons', [])
    cols = len(assigned_persons)

    if cols == 0:
        return None

    # Create footer table with two rows
    footer = doc.add_table(rows=2, cols=cols)  

    # Fill in the roles and names
    for i, person in enumerate(assigned_persons):
        cell = footer.rows[0].cells[i]
        paragraph = cell.paragraphs[0]
        
        # Add role (bold)
        role_run = paragraph.add_run(person.get('role', ''))
        role_run.bold = True
        role_run.font.size = Pt(11)
        
        # Add line break
        paragraph.add_run('\n')
        
        # Add name (not bold)
        name_run = paragraph.add_run(person.get('name', ''))
        name_run.font.size = Pt(11)
        
        # Center align the cell
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    footer_rows = footer.rows
    footer_rows[0].height = Inches(0.8)
    footer_rows[1].height = Inches(0.5)

    # Set no borders for all cells
    for row in footer.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                border_element = OxmlElement(f'w:{border}')
                border_element.set(qn('w:val'), 'nil')
                tcBorders.append(border_element)
            tcPr.append(tcBorders)

    # Add signature section
    merged_footer_row = footer.rows[1]
    merged_cells = merged_footer_row.cells[0].merge(merged_footer_row.cells[cols - 1])
    merged_cells.text = "Signature"
    merged_cells.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for run in merged_cells.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(14)
    
    return footer



# names = ["John Doe", "Jane Smith", "Bob Johnson"]

# # create document
# doc = create_document()
# create_header(doc)
# create_title_section(doc)
# create_object_section(doc)
# create_table_section(doc)
# create_footer(doc,3,names)




# # save document
# file_path = os.path.join(FILES_DIR, f"BreakdownSheet_4.docx")
# doc.save(file_path) 

if __name__ == "__main__":
    try:
        if len(sys.argv) < 2:
            print("Error: No data provided")
            sys.exit(1)
            
        # Parse the JSON data from command line argument
        data = json.loads(sys.argv[1])
        # print(data)

        #create document
        doc = create_document()
        create_header(doc)
        create_title_section(doc)
        create_object_section(doc)
        create_table_section(doc, data)
        create_footer(doc, data)

        #use the provided filename or generate a default one
        filename = f"BreakdownSheet_{data.get('fileName', 'FICHE PANNE')}.docx"
        file_path = os.path.join(FILES_DIR, filename)
        
        # Save the document
        doc.save(file_path)

        # Return JSON response
        print(json.dumps({
            "success": True,
            "message": "Document created successfully",
            "filepath": file_path,
            "filename": filename,
        }))

    #
    except Exception as e:
        print(json.dumps({
            "success": False,
            "message": str(e),
        }))
        sys.exit(1)

    
